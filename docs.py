import docx
doc = docx.Document()
doc.add_paragraph('testing')
doc.save('test.docx')
#doc.add_picture('image.jpg', width= docx.shared.Inches(4), height- docx.shared.Inches(3)) to add image to docx